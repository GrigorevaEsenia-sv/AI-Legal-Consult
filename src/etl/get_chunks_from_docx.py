import re
from typing import List, Dict, Tuple
from dataclasses import dataclass
from docx import Document
import hashlib
from datetime import datetime

@dataclass
class LegalChunk:
    section: str
    chapter: str
    article: str
    text: str
    metadata: Dict[str, str]
    chunk_id: str

def extract_document_metadata(file_path: str) -> Dict[str, str]:
    """Извлекает базовые метаданные из имени файла и пути"""
    return {
        "source": file_path,
        "document_type": "law",
        "jurisdiction": "RU",
        "file_name": file_path.split("/")[-1],
        "processing_date": datetime.now().isoformat()  # Теперь работает правильно
    }

def split_article_into_points(article_text: str) -> List[str]:
    """Разбивает текст статьи на пункты"""
    # Разбиваем по нумерованным пунктам (1., 2. и т.д.)
    points = re.split(r'\n(\d+\.)\s', article_text)
    
    # Объединяем разделители с текстом пунктов
    if len(points) > 1:
        processed_points = []
        for i in range(1, len(points), 2):
            point_num = points[i]
            point_text = points[i+1] if i+1 < len(points) else ""
            processed_points.append(f"{point_num} {point_text.strip()}")
        return processed_points
    
    # Если пунктов нет, возвращаем весь текст как один пункт
    return [article_text]

def split_large_point(point_text: str, metadata: Dict, max_size: int, overlap: int) -> List[LegalChunk]:
    """Разбивает большой пункт статьи на несколько чанков"""
    chunks = []
    words = point_text.split()
    current_chunk = []
    current_length = 0
    
    for word in words:
        current_chunk.append(word)
        current_length += len(word) + 1  # +1 за пробел
        
        if current_length >= max_size:
            chunk_text = ' '.join(current_chunk)
            chunks.append(create_chunk(
                metadata.get("section", ""),
                metadata.get("chapter", ""),
                metadata.get("article", ""),
                chunk_text,
                metadata.copy()
            ))
            
            # Сохраняем перекрытие
            current_chunk = current_chunk[-overlap:] if overlap else []
            current_length = sum(len(w) + 1 for w in current_chunk)
    
    if current_chunk:
        chunk_text = ' '.join(current_chunk)
        chunks.append(create_chunk(
            metadata.get("section", ""),
            metadata.get("chapter", ""),
            metadata.get("article", ""),
            chunk_text,
            metadata.copy()
        ))
    
    return chunks

def format_chunk_text(section: str, chapter: str, article: str, text: str) -> str:
    """Форматирует текст чанка с сохранением иерархии"""
    lines = []
    if section:
        lines.append(f"Раздел: {section}")
    if chapter:
        lines.append(f"Глава: {chapter}")
    if article:
        lines.append(f"Статья: {article}")
    lines.append("Текст:")
    lines.append(text)
    return "\n".join(lines)

def create_chunk(section: str, chapter: str, article: str, text: str, metadata: Dict) -> LegalChunk:
    """Создает объект LegalChunk с уникальным ID"""
    chunk_id = hashlib.md5(
        f"{section}_{chapter}_{article}_{text[:50]}".encode()
    ).hexdigest()
    
    return LegalChunk(
        section=section,
        chapter=chapter,
        article=article,
        text=text,
        metadata=metadata,
        chunk_id=chunk_id
    )

def identify_document_element(text: str) -> Tuple[str, str]:
    """Определяет тип структурного элемента документа"""
    # Регулярные выражения для определения структуры
    section_match = re.match(r'^Раздел\s*[IVXLCDM]+[.:]?\s*(.*)', text, re.IGNORECASE)
    chapter_match = re.match(r'^Глава\s*\d+[.:]?\s*(.*)', text, re.IGNORECASE)
    article_match = re.match(r'^Статья\s*\d+[.:]?\s*(.*)', text, re.IGNORECASE)
    
    if section_match:
        return ("SECTION", section_match.group(1).strip())
    elif chapter_match:
        return ("CHAPTER", chapter_match.group(1).strip())
    elif article_match:
        return ("ARTICLE", article_match.group(1).strip())
    return ("TEXT", text)


def process_legal_docx(file_path: str, max_chunk_size: int = 800, overlap: int = 100) -> List[LegalChunk]:
    """Обрабатывает юридический DOCX-документ"""
    doc = Document(file_path)
    current_section = "ОБЩИЕ ПОЛОЖЕНИЯ"
    current_chapter = ""
    current_article = ""
    chunks = []
    
    doc_metadata = extract_document_metadata(file_path)
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        element_type, element_content = identify_document_element(text)
        
        if element_type == "SECTION":
            current_section = element_content
        elif element_type == "CHAPTER":
            current_chapter = element_content
        elif element_type == "ARTICLE":
            current_article = element_content
            # Передаем все необходимые аргументы, включая overlap
            article_chunks = process_article(
                doc=doc,
                paragraph_idx=para_idx,
                section=current_section,
                chapter=current_chapter,
                article=current_article,
                metadata=doc_metadata,
                max_size=max_chunk_size,
                overlap=overlap  # Добавлен отсутствующий аргумент
            )
            chunks.extend(article_chunks)
    
    return chunks


def process_article(
    doc: Document,
    paragraph_idx: int,
    section: str,
    chapter: str,
    article: str,
    metadata: Dict[str, str],
    max_size: int,
    overlap: int  # Объявлен как обязательный параметр
) -> List[LegalChunk]:
    """Обрабатывает статью и разбивает на чанки"""
    article_text = [doc.paragraphs[paragraph_idx].text]
    
    # Собираем все параграфы статьи (до следующего заголовка)
    for next_para in doc.paragraphs[paragraph_idx+1:]:
        next_para_text = next_para.text.strip()
        
        # Проверяем, не является ли параграф новым заголовком
        is_heading = (
            re.match(r'^(Раздел|Глава|Статья)\s', next_para_text) or 
            (hasattr(next_para, 'style') and next_para.style.name.lower().startswith('heading'))
        )
        
        if is_heading:
            break
            
        article_text.append(next_para_text)
    
    full_text = "\n".join(article_text)
    points = split_article_into_points(full_text)
    chunks = []
    
    for point_num, point_text in enumerate(points, 1):
        chunk_text = format_chunk_text(section, chapter, article, point_text)
        point_metadata = metadata.copy()
        point_metadata.update({
            "point_number": point_num,
            "point_text": point_text[:100] + "..." if len(point_text) > 100 else point_text
        })
        
        if len(chunk_text) > max_size:
            sub_chunks = split_large_point(chunk_text, point_metadata, max_size, overlap)
            chunks.extend(sub_chunks)
        else:
            chunks.append(create_chunk(
                section, chapter, article, 
                chunk_text, point_metadata
            ))
    
    return chunks

def process_article(doc, paragraph_idx, section, chapter, article, metadata, max_size, overlap) -> List[LegalChunk]:
    """Обрабатывает статью и разбивает на чанки"""
    article_text = [doc.paragraphs[paragraph_idx].text]
    
    # Собираем все параграфы статьи (до следующего заголовка)
    for next_para in doc.paragraphs[paragraph_idx+1:]:
        next_para_text = next_para.text.strip()
        
        # Проверяем, не является ли параграф новым заголовком
        is_heading = (
            re.match(r'^(Раздел|Глава|Статья)\s', next_para_text) or 
            next_para.style.name.lower().startswith('heading')
        )
        
        if is_heading:
            break
            
        article_text.append(next_para_text)
    
    full_text = "\n".join(article_text)
    
    # Разбиваем статью на пункты
    points = split_article_into_points(full_text)
    
    # Формируем чанки для каждого пункта
    chunks = []
    for point_num, point_text in enumerate(points, 1):
        # Формируем текст чанка с полной структурой
        chunk_text = format_chunk_text(section, chapter, article, point_text)
        
        # Создаем метаданные с информацией о пункте
        point_metadata = metadata.copy()
        point_metadata.update({
            "point_number": point_num,
            "point_text": point_text[:100] + "..." if len(point_text) > 100 else point_text
        })
        
        # Если пункт слишком большой, разбиваем его
        if len(chunk_text) > max_size:
            sub_chunks = split_large_point(chunk_text, point_metadata, max_size, overlap)
            chunks.extend(sub_chunks)
        else:
            chunks.append(create_chunk(
                section, chapter, article, 
                chunk_text, point_metadata
            ))
    
    return chunks


# Пример использования
if __name__ == "__main__":
    doc_path = 'data/raw/housing_code/JKRF.docx'
    chunks = process_legal_docx(doc_path)
    for chunk in chunks[:7]:
        print(chunk)
        print('='*50)